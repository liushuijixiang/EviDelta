from __future__ import annotations

import csv
from datetime import datetime
import json

from openpyxl import Workbook
from openpyxl.workbook.defined_name import DefinedName
from docx import Document

from feishu_agent_bot.parsers.csv_parser import CsvParser
from feishu_agent_bot.parsers.excel_parser import ExcelParser
from feishu_agent_bot.parsers.html_parser import HtmlParser
from feishu_agent_bot.parsers.structured import DocxParser, JsonParser


def test_csv_parser_detects_bom_delimiter_dates_missing_and_thousands(tmp_path):
    path = tmp_path / "metrics.csv"
    path.write_text(
        '名称;金额;日期;备注\nA;"1,234.50";2026/06/19;\n',
        encoding="utf-8-sig",
    )

    parsed = CsvParser(chunk_size=1).parse(path, asset_id="CSV1")

    assert parsed.metadata["encoding"] == "utf-8-sig"
    assert parsed.metadata["delimiter"] == ";"
    assert parsed.tables[0].rows == [
        {"名称": "A", "金额": 1234.5, "日期": "2026-06-19", "备注": None}
    ]
    assert (tmp_path / "metrics.csv.normalized.csv").is_file()


def test_csv_parser_supports_gb18030_tsv_comments_and_duplicate_headers(tmp_path):
    path = tmp_path / "prices.tsv"
    path.write_bytes(
        "# 导出备注\n产品\t价格\t价格\n甲\t99\t109\n".encode("gb18030")
    )

    parsed = CsvParser(chunk_size=1).parse(path, asset_id="CSV2")

    assert parsed.metadata["encoding"] == "gb18030"
    assert parsed.metadata["delimiter"] == "\t"
    assert parsed.tables[0].columns == ["产品", "价格", "价格_2"]
    assert parsed.tables[0].rows[0] == {"产品": "甲", "价格": 99, "价格_2": 109}
    assert any("duplicate_column=价格" in item for item in parsed.warnings)


def test_csv_parser_chunks_full_normalized_file_and_limits_analysis_rows(tmp_path):
    path = tmp_path / "large.csv"
    path.write_text("id,value\n1,10\n2,20\n3,30\n4,40\n", encoding="utf-8")

    parsed = CsvParser(max_rows=2, preview_rows=1, chunk_size=2).parse(
        path, asset_id="CSV3"
    )

    assert parsed.metadata["row_count"] == 4
    assert parsed.metadata["analysis_row_count"] == 2
    assert parsed.metadata["truncated_for_analysis"] is True
    assert len(parsed.tables[0].rows) == 2
    with open(parsed.metadata["normalized_path"], encoding="utf-8", newline="") as file:
        assert len(list(csv.DictReader(file))) == 4


def test_excel_parser_preserves_sheet_formula_format_and_lineage_metadata(tmp_path):
    path = tmp_path / "benchmark.xlsx"
    workbook = Workbook()
    sheet = workbook.active
    sheet.title = "报价 表"
    sheet.merge_cells("A1:A2")
    sheet.merge_cells("B1:C1")
    sheet.merge_cells("D1:D2")
    sheet.merge_cells("E1:E2")
    sheet["A1"] = "产品"
    sheet["B1"] = "价格"
    sheet["B2"] = "月费"
    sheet["C2"] = "年费"
    sheet["D1"] = "增长率"
    sheet["E1"] = "总价"
    sheet.append(["标准版", 99, 999, 0.25, "=B3+C3"])
    sheet["B3"].number_format = '¥#,##0.00'
    sheet["D3"].number_format = "0.0%"
    sheet["F1"] = "更新日期"
    sheet["F2"] = "日期"
    sheet["F3"] = datetime(2026, 6, 19)
    hidden = workbook.create_sheet("内部")
    hidden.sheet_state = "hidden"
    hidden.append(["键", "值"])
    hidden.append(["成本", 42])
    workbook.defined_names.add(
        DefinedName("PriceArea", attr_text="'报价 表'!$B$3:$C$3")
    )
    workbook.save(path)

    parsed = ExcelParser().parse(path, asset_id="XLSX1")

    assert parsed.metadata["sheet_count"] == 2
    assert parsed.metadata["macros_loaded"] is False
    assert parsed.metadata["external_links_loaded"] is False
    quote_table = parsed.tables[0]
    assert quote_table.columns[:5] == [
        "产品",
        "价格 / 月费",
        "价格 / 年费",
        "增长率",
        "总价",
    ]
    assert quote_table.rows[0]["价格 / 月费"] == 99
    assert quote_table.source_locator.endswith(
        "#sheet=%E6%8A%A5%E4%BB%B7%20%E8%A1%A8&range=A1:F3"
    )
    assert parsed.tables[1].metadata["sheet_state"] == "hidden"
    formulas = [
        item
        for item in quote_table.metadata["cell_metadata"]
        if item["formula"]
    ]
    assert formulas == [
        {
            "asset_id": "XLSX1",
            "sheet_name": "报价 表",
            "cell_range": "E3",
            "column": "总价",
            "formula": "=B3+C3",
            "cached_value": None,
            "calculation_status": "stale_or_unknown",
            "number_format": "General",
        }
    ]
    assert any(item["name"] == "PriceArea" for item in parsed.metadata["named_ranges"])
    assert any("stale_or_unknown" in item for item in parsed.warnings)
    formats = {
        item["cell_range"]: item["number_format"]
        for item in quote_table.metadata["cell_metadata"]
    }
    assert formats["B3"] == '¥#,##0.00'
    assert formats["D3"] == "0.0%"
    assert formats["F3"] == "yyyy-mm-dd h:mm:ss"


def test_html_parser_expands_spans_and_links_table_to_context(tmp_path):
    path = tmp_path / "pricing.html"
    path.write_text(
        """
        <html><head><title>产品报告</title></head><body>
          <h2>套餐比较</h2><p>以下价格来自官方产品页。</p>
          <table><caption>公开报价</caption>
            <tr><th rowspan="2">产品</th><th colspan="2">价格</th></tr>
            <tr><th>月费</th><th>年费</th></tr>
            <tr><td>标准版</td><td>99</td><td>999</td></tr>
          </table>
        </body></html>
        """,
        encoding="utf-8",
    )

    parsed = HtmlParser().parse(path, asset_id="HTML1")

    table = parsed.tables[0]
    assert parsed.title == "产品报告"
    assert table.caption == "公开报价"
    assert table.columns == ["产品", "价格 / 月费", "价格 / 年费"]
    assert table.rows[0] == {
        "产品": "标准版",
        "价格 / 月费": 99,
        "价格 / 年费": 999,
    }
    assert table.metadata["context"] == {
        "preceding_text": "以下价格来自官方产品页。",
        "heading": "套餐比较",
    }
    assert table.metadata["rowspan_colspan_expanded"] is True
    assert table.source_locator == "pricing.html#table-1"


def test_json_parser_discovers_nested_api_tables_and_limits_text(tmp_path):
    path = tmp_path / "metrics.json"
    path.write_text(
        json.dumps(
            {
                "status": "ok",
                "data": {
                    "series": [
                        {
                            "month": "2026-01",
                            "metrics": {"revenue": "¥1,000", "growth_pct": "5%"},
                        },
                        {
                            "month": "2026-02",
                            "metrics": {"revenue": "¥1,200", "growth_pct": "6%"},
                        },
                    ],
                    "items": [{"name": "A", "score": 1}],
                },
            },
            ensure_ascii=False,
        ),
        encoding="utf-8",
    )

    parsed = JsonParser(max_text_length=40).parse(path, asset_id="JSON1")

    assert parsed.metadata["table_count"] == 2
    assert parsed.tables[0].caption == "$.data.series"
    assert parsed.tables[0].rows[0] == {
        "month": "2026-01",
        "metrics.revenue": "¥1,000",
        "metrics.growth_pct": "5%",
    }
    assert parsed.tables[0].source_locator == "metrics.json#jsonpath=$.data.series"
    assert any("json_text_truncated=40" in item for item in parsed.warnings)


def test_docx_parser_extracts_headings_properties_tables_and_notes(tmp_path):
    path = tmp_path / "industry.docx"
    document = Document()
    document.core_properties.title = "行业报告"
    document.core_properties.author = "Research Team"
    document.add_heading("市场概览", level=1)
    document.add_paragraph("正文段落。")
    table = document.add_table(rows=2, cols=2)
    table.rows[0].cells[0].text = "公司"
    table.rows[0].cells[1].text = "收入"
    table.rows[1].cells[0].text = "A"
    table.rows[1].cells[1].text = "1,000"
    document.save(path)

    parsed = DocxParser().parse(path, asset_id="DOCX1")

    assert parsed.title == "行业报告"
    assert parsed.metadata["author"] == "Research Team"
    assert parsed.metadata["macros_executed"] is False
    assert parsed.text_blocks[0].section == "市场概览"
    assert parsed.text_blocks[1].section == "市场概览"
    assert parsed.tables[0].rows[0] == {"公司": "A", "收入": 1000}
    assert parsed.tables[0].source_locator == "industry.docx#table-1"
