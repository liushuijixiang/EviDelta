from __future__ import annotations

import json
from pathlib import Path
import zipfile
from xml.etree import ElementTree

from .base import ParsedAsset, ParsedTable, TextBlock
from .csv_parser import CsvParser
from .excel_parser import ExcelParser
from .html_parser import HtmlParser
from .tabular_utils import normalize_scalar, unique_headers


def _records_to_table(
    table_id: str,
    rows: list[dict[str, object]],
    *,
    caption: str | None = None,
    source_locator: str | None = None,
    extraction_method: str = "json_table",
    metadata: dict[str, object] | None = None,
) -> ParsedTable:
    columns: list[str] = []
    for row in rows:
        for key in row:
            if key not in columns:
                columns.append(str(key))
    return ParsedTable(
        table_id=table_id,
        columns=columns,
        rows=rows,
        caption=caption,
        source_locator=source_locator,
        extraction_method=extraction_method,
        metadata=metadata or {},
    )


class JsonParser:
    name = "json"
    version = "2.0"
    supported_file_types = {"json"}
    supported_mime_types = {"application/json", "application/ld+json"}

    def __init__(
        self,
        *,
        max_depth: int = 8,
        max_objects: int = 100_000,
        max_fields: int = 200,
        max_text_length: int = 2_000_000,
    ) -> None:
        self.max_depth = max_depth
        self.max_objects = max_objects
        self.max_fields = max_fields
        self.max_text_length = max_text_length

    def can_parse(self, asset) -> bool:
        return asset.file_type in self.supported_file_types or asset.detected_mime_type in self.supported_mime_types

    def parse(self, path: Path, *, asset_id: str) -> ParsedAsset:
        data = json.loads(path.read_text(encoding="utf-8"))
        warnings: list[str] = []
        tables: list[ParsedTable] = []
        object_counter = {"count": 0}
        candidates = self._find_tables(data, "$", 0, object_counter, warnings)
        for index, (json_path, rows) in enumerate(candidates, start=1):
            flat_rows = [self._flatten_object(row, warnings=warnings) for row in rows]
            if not flat_rows:
                continue
            if len(flat_rows[0]) > self.max_fields:
                warnings.append(f"json_table_fields_truncated={json_path}; max_fields={self.max_fields}")
                flat_rows = [
                    dict(list(row.items())[: self.max_fields]) for row in flat_rows
                ]
            tables.append(
                _records_to_table(
                    f"{asset_id}-T{index:03d}",
                    flat_rows,
                    caption=json_path,
                    source_locator=f"{path.name}#jsonpath={json_path}",
                    metadata={
                        "json_path": json_path,
                        "row_count": len(flat_rows),
                        "max_depth": self.max_depth,
                    },
                )
            )
        text = json.dumps(data, ensure_ascii=False, indent=2, default=str)
        if len(text) > self.max_text_length:
            text = text[: self.max_text_length]
            warnings.append(f"json_text_truncated={self.max_text_length}")
        return ParsedAsset(
            asset_id,
            "json",
            None,
            [TextBlock(f"{asset_id}-B001", text, source_locator=f"{path.name}#json")],
            tables,
            metadata={
                "table_count": len(tables),
                "object_count": object_counter["count"],
                "max_depth": self.max_depth,
                "max_objects": self.max_objects,
                "max_fields": self.max_fields,
                "max_text_length": self.max_text_length,
            },
            warnings=warnings,
            extraction_method="json",
        )

    def _find_tables(
        self,
        value: object,
        path: str,
        depth: int,
        counter: dict[str, int],
        warnings: list[str],
    ) -> list[tuple[str, list[dict[str, object]]]]:
        if depth > self.max_depth:
            warnings.append(f"json_max_depth_reached={path}")
            return []
        if isinstance(value, dict):
            counter["count"] += 1
            if counter["count"] > self.max_objects:
                warnings.append(f"json_max_objects_reached={self.max_objects}")
                return []
            tables: list[tuple[str, list[dict[str, object]]]] = []
            for key, item in value.items():
                tables.extend(
                    self._find_tables(item, f"{path}.{key}", depth + 1, counter, warnings)
                )
            return tables
        if isinstance(value, list):
            if self._is_records(value):
                return [(path, value)]  # type: ignore[list-item]
            tables = []
            for index, item in enumerate(value[: self.max_objects]):
                tables.extend(
                    self._find_tables(item, f"{path}[{index}]", depth + 1, counter, warnings)
                )
            return tables
        return []

    @staticmethod
    def _is_records(value: list[object]) -> bool:
        if not value or not all(isinstance(item, dict) for item in value):
            return False
        dicts = [item for item in value if isinstance(item, dict)]
        if not dicts:
            return False
        scalar_counts = [
            sum(not isinstance(field, (dict, list)) for field in row.values())
            for row in dicts[:20]
        ]
        return bool(scalar_counts) and max(scalar_counts) > 0

    def _flatten_object(
        self,
        value: dict[str, object],
        *,
        prefix: str = "",
        warnings: list[str],
        depth: int = 0,
    ) -> dict[str, object]:
        if depth > self.max_depth:
            warnings.append(f"json_flatten_depth_reached={prefix or '$'}")
            return {}
        result: dict[str, object] = {}
        for key, item in value.items():
            name = f"{prefix}.{key}" if prefix else str(key)
            if isinstance(item, dict):
                result.update(
                    self._flatten_object(item, prefix=name, warnings=warnings, depth=depth + 1)
                )
            elif isinstance(item, list):
                result[name] = json.dumps(item, ensure_ascii=False, default=str)
            else:
                result[name] = normalize_scalar(item)
        return result


class DocxParser:
    name = "docx"
    version = "2.0"
    supported_file_types = {"docx"}
    supported_mime_types = {
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    }

    def can_parse(self, asset) -> bool:
        return asset.file_type in self.supported_file_types or asset.detected_mime_type in self.supported_mime_types

    def parse(self, path: Path, *, asset_id: str) -> ParsedAsset:
        from docx import Document

        document = Document(path)
        blocks: list[TextBlock] = []
        heading_stack: list[str] = []
        block_index = 1
        for paragraph_index, paragraph in enumerate(document.paragraphs, start=1):
            text = paragraph.text.strip()
            if not text:
                continue
            style = paragraph.style.name if paragraph.style is not None else ""
            if style.lower().startswith("heading") or style.startswith("标题"):
                level = self._heading_level(style)
                heading_stack = heading_stack[: max(0, level - 1)] + [text]
                section = " / ".join(heading_stack)
            else:
                section = " / ".join(heading_stack) if heading_stack else None
            blocks.append(
                TextBlock(
                    f"{asset_id}-B{block_index:03d}",
                    text,
                    section=section,
                    source_locator=f"{path.name}#paragraph-{paragraph_index}",
                )
            )
            block_index += 1
        extra_blocks = self._extract_docx_notes(path, asset_id, start_index=block_index)
        blocks.extend(extra_blocks)
        tables: list[ParsedTable] = []
        for index, table in enumerate(document.tables, start=1):
            matrix = [[cell.text.strip() for cell in row.cells] for row in table.rows]
            if not matrix:
                continue
            headers, header_warnings = unique_headers(matrix[0])
            rows = [
                {
                    headers[i] if i < len(headers) else f"column_{i + 1}": normalize_scalar(value)
                    for i, value in enumerate(row)
                }
                for row in matrix[1:]
            ]
            tables.append(
                ParsedTable(
                    f"{asset_id}-T{index:03d}",
                    headers,
                    rows,
                    source_locator=f"{path.name}#table-{index}",
                    extraction_method="docx_table",
                    metadata={"warnings": header_warnings},
                )
            )
        properties = document.core_properties
        metadata = {
            "title": properties.title,
            "author": properties.author,
            "subject": properties.subject,
            "keywords": properties.keywords,
            "created": properties.created.isoformat() if properties.created else None,
            "modified": properties.modified.isoformat() if properties.modified else None,
            "paragraph_count": len(blocks),
            "table_count": len(tables),
            "macros_executed": False,
            "embedded_objects_executed": False,
        }
        return ParsedAsset(
            asset_id,
            "docx",
            properties.title or None,
            blocks,
            tables,
            metadata=metadata,
            extraction_method="docx",
        )

    @staticmethod
    def _heading_level(style_name: str) -> int:
        for token in reversed(style_name.split()):
            if token.isdigit():
                return max(1, int(token))
        return 1

    @staticmethod
    def _extract_docx_notes(path: Path, asset_id: str, *, start_index: int) -> list[TextBlock]:
        blocks: list[TextBlock] = []
        namespaces = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
        targets = [
            ("word/footnotes.xml", "footnote"),
            ("word/comments.xml", "comment"),
        ]
        try:
            with zipfile.ZipFile(path) as archive:
                for name, label in targets:
                    if name not in archive.namelist():
                        continue
                    root = ElementTree.fromstring(archive.read(name))
                    for node in root:
                        texts = [
                            item.text or ""
                            for item in node.findall(".//w:t", namespaces)
                            if item.text
                        ]
                        text = "".join(texts).strip()
                        if text:
                            blocks.append(
                                TextBlock(
                                    f"{asset_id}-B{start_index + len(blocks):03d}",
                                    text,
                                    section=label,
                                    source_locator=f"{path.name}#{label}-{len(blocks) + 1}",
                                )
                            )
        except (OSError, zipfile.BadZipFile, ElementTree.ParseError):
            return blocks
        return blocks


class TextParser:
    name = "text"
    version = "1.0"
    supported_file_types = {"text"}
    supported_mime_types = {"text/plain"}

    def can_parse(self, asset) -> bool:
        return asset.file_type in self.supported_file_types or asset.detected_mime_type in self.supported_mime_types

    def parse(self, path: Path, *, asset_id: str) -> ParsedAsset:
        text = path.read_text(encoding="utf-8", errors="ignore")
        return ParsedAsset(asset_id, "text", None, [TextBlock(f"{asset_id}-B001", text, source_locator=f"{path.name}#text")], extraction_method="text")
