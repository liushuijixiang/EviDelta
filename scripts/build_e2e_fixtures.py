from __future__ import annotations

import json
from pathlib import Path

from docx import Document
import fitz
from openpyxl import Workbook
from openpyxl.workbook.defined_name import DefinedName


ROOT = Path(__file__).resolve().parents[1]
FIXTURE_DIR = ROOT / "tests" / "fixtures"


def main() -> int:
    FIXTURE_DIR.mkdir(parents=True, exist_ok=True)
    build_text_pdf(FIXTURE_DIR / "sample_text.pdf")
    build_scanned_pdf(FIXTURE_DIR / "sample_scanned.pdf")
    build_prices_csv(FIXTURE_DIR / "sample_prices.csv")
    build_competitors_xlsx(FIXTURE_DIR / "sample_competitors.xlsx")
    build_metrics_json(FIXTURE_DIR / "sample_metrics.json")
    build_report_docx(FIXTURE_DIR / "sample_report.docx")
    build_page_html(FIXTURE_DIR / "sample_page.html")
    return 0


def build_text_pdf(path: Path) -> None:
    document = fitz.open()
    page = document.new_page()
    page.insert_text((72, 72), "Competitor A supports fast charging and offline mode.")
    page.insert_text((72, 100), "Competitor B focuses on low monthly price.")
    document.set_metadata({"title": "Fixture Product Features"})
    document.save(path)
    document.close()


def build_scanned_pdf(path: Path) -> None:
    document = fitz.open()
    page = document.new_page()
    pixmap = fitz.Pixmap(fitz.csRGB, (0, 0, 80, 40), False)
    pixmap.clear_with(255)
    page.insert_image(fitz.Rect(72, 72, 220, 150), pixmap=pixmap)
    document.set_metadata({"title": "Fixture Scanned Page"})
    document.save(path)
    document.close()


def build_prices_csv(path: Path) -> None:
    path.write_text(
        "\n".join(
            [
                "month,company,price_cny,sales_units",
                "2026-01,Competitor A,99,1200",
                "2026-02,Competitor A,109,1350",
                "2026-01,Competitor B,89,980",
                "2026-02,Competitor B,95,1010",
            ]
        )
        + "\n",
        encoding="utf-8",
    )


def build_competitors_xlsx(path: Path) -> None:
    workbook = Workbook()
    sheet = workbook.active
    sheet.title = "Competitors"
    sheet.merge_cells("A1:A2")
    sheet.merge_cells("B1:C1")
    sheet.merge_cells("D1:D2")
    sheet["A1"] = "company"
    sheet["B1"] = "plan"
    sheet["B2"] = "tier"
    sheet["C2"] = "monthly_price_cny"
    sheet["D1"] = "feature"
    sheet.append(["Competitor A", "Pro", 109, "fast charging"])
    sheet.append(["Competitor B", "Basic", 95, "offline mode"])
    sheet["C3"].number_format = '¥#,##0'
    sheet["C4"].number_format = '¥#,##0'
    hidden = workbook.create_sheet("Internal")
    hidden.sheet_state = "hidden"
    hidden.append(["note", "fixture hidden sheet"])
    workbook.defined_names.add(
        DefinedName("CompetitorRange", attr_text="'Competitors'!$A$3:$D$4")
    )
    workbook.save(path)


def build_metrics_json(path: Path) -> None:
    path.write_text(
        json.dumps(
            {
                "api_version": "fixture-v1",
                "data": {
                    "time_series": [
                        {"month": "2026-01", "active_users": 10000, "growth_pct": 5},
                        {"month": "2026-02", "active_users": 11200, "growth_pct": 12},
                    ]
                },
            },
            ensure_ascii=False,
            indent=2,
        ),
        encoding="utf-8",
    )


def build_report_docx(path: Path) -> None:
    document = Document()
    document.core_properties.title = "Fixture Industry Report"
    document.core_properties.author = "Fixture Research"
    document.add_heading("行业报告摘要", level=1)
    document.add_paragraph("本行业报告指出，中端套餐增长来自企业移动办公场景。")
    document.add_heading("竞争格局", level=2)
    document.add_paragraph("Competitor A 在快充功能上领先，Competitor B 强调低价套餐。")
    table = document.add_table(rows=3, cols=2)
    table.rows[0].cells[0].text = "company"
    table.rows[0].cells[1].text = "positioning"
    table.rows[1].cells[0].text = "Competitor A"
    table.rows[1].cells[1].text = "feature leader"
    table.rows[2].cells[0].text = "Competitor B"
    table.rows[2].cells[1].text = "price challenger"
    document.save(path)


def build_page_html(path: Path) -> None:
    path.write_text(
        """
        <html><head><title>Official Fixture Product Page</title></head>
        <body>
          <h1>官方产品页</h1>
          <p>该页面列出公开套餐能力和价格。</p>
          <table><caption>官方套餐</caption>
            <tr><th rowspan="2">company</th><th colspan="2">pricing</th><th rowspan="2">source</th></tr>
            <tr><th>monthly_cny</th><th>annual_cny</th></tr>
            <tr><td>Competitor A</td><td>109</td><td>1090</td><td>official</td></tr>
            <tr><td>Competitor B</td><td>95</td><td>950</td><td>official</td></tr>
          </table>
        </body></html>
        """.strip(),
        encoding="utf-8",
    )


if __name__ == "__main__":
    raise SystemExit(main())
